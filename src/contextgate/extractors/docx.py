from pathlib import Path

import docx
from docx.oxml.ns import qn

from ..exceptions import ExtractionError, FileTooLargeError
from ..result import ExtractedSegment

MAX_FILE_BYTES = 20 * 1024 * 1024


def _check_size(path: Path) -> None:
    if path.stat().st_size > MAX_FILE_BYTES:
        raise FileTooLargeError(f"File exceeds 20MB: {path}")


def _paragraph_text(para) -> str:
    return para.text


def _table_text(table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def _section_text(section_part) -> list[str]:
    texts: list[str] = []
    if section_part is None:
        return texts
    for para in section_part.paragraphs:
        t = _paragraph_text(para)
        if t.strip():
            texts.append(t)
    return texts


def extract_docx(path: str | Path) -> list[ExtractedSegment]:
    p = Path(path)
    _check_size(p)
    try:
        doc = docx.Document(str(p))
    except Exception as e:
        raise ExtractionError(f"Failed to read DOCX: {p}: {e}") from e

    parts: list[str] = []

    # headers and footers
    for section in doc.sections:
        parts.extend(_section_text(section.header))
        parts.extend(_section_text(section.footer))

    # paragraphs
    for para in doc.paragraphs:
        t = _paragraph_text(para)
        if t.strip():
            parts.append(t)

    # tables
    for table in doc.tables:
        t = _table_text(table)
        if t.strip():
            parts.append(t)

    combined = "\n".join(parts)
    if not combined.strip():
        return [ExtractedSegment(text="")]
    return [ExtractedSegment(text=combined)]
